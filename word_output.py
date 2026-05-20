"""
word_output.py
Structure per table:
  [Heading]   Title from text above the table in the PDF
  [Bullets]   AI-generated summary points
  [Snapshot]  Full table screenshot as-is
  [Divider]
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

from pdf_table_extractor import snapshot_word_width


def build_word_report(
    tables: list,
    pdf_filename: str = None,
    ref_filename: str = None,
) -> bytes:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    for i, table in enumerate(tables, start=1):
        # Priority: match_label (ref cross-match) → first header cell → fallback
        raw_label = (table.get("match_label") or _derive_label(table, i))
        label     = raw_label.split(":")[0].strip()
        summary  = table.get("summary", ["No summary available."])
        snapshot = table.get("snapshot")

        # Heading
        h = doc.add_heading(label, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Bullets
        for bullet in summary:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(bullet).font.size = Pt(10)

        doc.add_paragraph()

        # Snapshot
        if snapshot:
            buf = io.BytesIO()
            snapshot.save(buf, format="PNG")
            buf.seek(0)
            doc.add_picture(buf, width=Inches(snapshot_word_width(snapshot)))
        else:
            ph = doc.add_paragraph("[Table snapshot not available]")
            ph.runs[0].font.italic    = True
            ph.runs[0].font.size      = Pt(9)
            ph.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        doc.add_paragraph()
        _add_rule(doc)
        doc.add_paragraph()

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def _derive_label(table: dict, index: int) -> str:
    headers = [str(h).strip() for h in table.get("headers", []) if str(h).strip()]
    return headers[0] if headers else f"Table {index}" 


def _add_rule(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)
